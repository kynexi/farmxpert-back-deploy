import os
import re
from flask import send_file
import io
import requests
from flask import Blueprint, request, jsonify, send_file
from docx import Document
from bson import ObjectId
from app.db_utilis import client

import tempfile
from docx import Document
from fpdf import FPDF  

bp = Blueprint("main", __name__)
scraper_bp = Blueprint("scraper", __name__, url_prefix="/api/")

# helpers 

NBSP = "\u00A0"
BLANKISH_LINE = re.compile(r"^[\s\.\-–—_·•]{3,}$")
LABEL_LINE = re.compile(r"^[A-Za-zĂÂÎȘȚăâîșț0-9/(),.\- ]{3,120}$")
_BLANK_RE = re.compile(r"(?:_{4,}|\.{4,}|_{2,}\.{2,}|\.{2,}_{2,})")

def _clean_ws(s: str) -> str:
    return (s or "").replace(NBSP, " ").strip()

def _is_blankish(s: str) -> bool:
    t = _clean_ws(s)
    if not t:
        return True
    if BLANKISH_LINE.match(t):
        return True
    return t in {".", "..", "…", "-"}

def _is_labelish(text: str) -> bool:
    t = (text or "").strip()
    if not t or len(t) > 120:
        return False
    if not LABEL_LINE.match(t):
        return False
    if re.search(r"\b\d{6,}\b", t):
        return False
    return True

def _slug(s: str) -> str:
    s = re.sub(r"\s+", "_", (s or "").strip().lower())
    s = re.sub(r"[^a-z0-9_]", "", s)
    return s or "field"

def _guess_label(text: str, match: re.Match) -> str:
    left = text[:match.start()]
    m = re.search(r"\(([^\)]+)\)\s*$", left)
    if m:
        return m.group(1).strip()
    m = re.search(r"([A-Za-zăâîșțA-Z0-9/ ,\-]{3,})[: ]\s*$", left)
    if m:
        return m.group(1).strip()
    tokens = re.findall(r"[A-Za-zăâîșț0-9]+", left)[-4:]
    return " ".join(tokens).strip() or "câmp"

def _extract_fields_from_paragraph_text(text: str):
    fields = []
    for m in _BLANK_RE.finditer(text):
        label = _guess_label(text, m)
        context = text[max(0, m.start()-80): m.end()+40]
        fields.append({"id": _slug(label), "label": label, "context": context})
    return fields

def _extract_doc_fields(doc: Document):
    fields = []
    for p in doc.paragraphs:
        fields.extend(_extract_fields_from_paragraph_text(p.text))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                fields.extend(_extract_fields_from_paragraph_text(cell.text))
    seen = set()
    uniq = []
    for f in fields:
        k = f["id"]
        if k in seen: 
            continue
        seen.add(k)
        uniq.append(f)
    return uniq

def _flatten_profile(profile: dict) -> dict:
    out = {}
    if not isinstance(profile, dict):
        return out

    def add(k, v):
        k = (_slug(k).replace("_", " ") or "").strip().lower()
        if not k:
            return
        out[k] = str(v)

    def walk(prefix, d):
        if isinstance(d, dict):
            for k, v in d.items():
                key = f"{prefix} {k}".strip()
                if isinstance(v, (dict, list)):
                    walk(key, v)
                else:
                    add(key, v)
        elif isinstance(d, list):
            for i, v in enumerate(d, 1):
                walk(f"{prefix} {i}", v)
        else:
            add(prefix, d)

    walk("", profile)

    user = profile.get("user", {}) or {}
    fn = user.get("firstName") or user.get("first_name")
    ln = user.get("lastName") or user.get("last_name")
    if fn or ln:
        full = f"{fn or ''} {ln or ''}".strip()
        out["numele si prenumele"] = full
        out["numele și prenumele"] = full

    if user.get("phone"):
        out["telefon"] = str(user["phone"])
    if user.get("email"):
        out["email"] = str(user["email"])
        out["e mail"] = str(user["email"])

    return out

def _jaccard(a: str, b: str) -> float:
    ta = set(re.findall(r"[a-z0-9ăâîșț]+", (a or "").lower()))
    tb = set(re.findall(r"[a-z0-9ăâîșț]+", (b or "").lower()))
    return 0.0 if not ta or not tb else len(ta & tb) / len(ta | tb)

def _best_suggestion_for(label: str, suggestions: dict):
    best_k, best_s = None, 0.0
    for k in suggestions.keys():
        s = _jaccard(label, k)
        if s > best_s:
            best_k, best_s = k, s
    return best_k

def _copy_run_formatting(src_run, dst_run):
    try:
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
    except Exception:
        pass
    try:
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
    except Exception:
        pass
    try:
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
    except Exception:
        pass
    try:
        if src_run.font.color and src_run.font.color.rgb:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass

def _apply_suggestions_inline(doc: Document, suggestions: dict):
    def replace_text(txt: str):
        out, idx = [], 0
        for m in _BLANK_RE.finditer(txt):
            out.append(txt[idx:m.end()])
            label = _guess_label(txt, m)
            key = _best_suggestion_for(_slug(label), suggestions) or _best_suggestion_for(label, suggestions)
            val = suggestions.get(key) if key else None
            if val:
                out.append(f" «{val}»")
            idx = m.end()
        out.append(txt[idx:])
        return "".join(out)

    for p in list(doc.paragraphs):
        original = p.text or ""
        new_text = replace_text(original)
        if new_text != original:
            first_run = p.runs[0] if p.runs else None
            for run in list(p.runs):
                try:
                    run.text = ""
                except Exception:
                    pass
            new_run = p.add_run(new_text)
            if first_run:
                try:
                    _copy_run_formatting(first_run, new_run)
                except Exception:
                    pass

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if not cell.paragraphs:
                    continue
                cell_text = cell.text or ""
                new_text = replace_text(cell_text)
                if new_text != cell_text:
                    first_para = cell.paragraphs[0]
                    first_run = first_para.runs[0] if first_para.runs else None
                    for para in list(cell.paragraphs):
                        for run in list(para.runs):
                            try:
                                run.text = ""
                            except Exception:
                                pass
                        try:
                            para.text = ""
                        except Exception:
                            pass
                    target_para = cell.paragraphs[0]
                    try:
                        for rr in list(target_para.runs):
                            rr.text = ""
                    except Exception:
                        pass
                    new_run = target_para.add_run(new_text)
                    if first_run:
                        try:
                            _copy_run_formatting(first_run, new_run)
                        except Exception:
                            pass

def _apply_label_and_cell_fill(doc: Document, suggestions: dict):
    def pick(label: str):
        lbln = _slug(label).replace("_", " ")
        if lbln in suggestions:
            return suggestions.get(lbln)
        key = _best_suggestion_for(_slug(label), suggestions) or _best_suggestion_for(label, suggestions)
        return suggestions.get(key)

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            for i in range(len(cells) - 1):
                left_txt = _clean_ws(cells[i].text)
                right_txt = cells[i+1].text
                if _is_labelish(left_txt) and _is_blankish(right_txt):
                    val = pick(left_txt)
                    if val:
                        cells[i+1].text = val

    paras = list(doc.paragraphs)
    for idx, p in enumerate(paras):
        t = _clean_ws(p.text)
        if not _is_labelish(t):
            continue
        nxt = paras[idx+1].text if idx + 1 < len(paras) else ""
        if not _is_blankish(nxt):
            continue
        val = pick(t)
        if val:
            p.add_run(f" «{val}»")

def _ai_fill_fields(fields, profile=None, instructions=None, language="ro"):
    try:
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            from openai import OpenAI
            import json
            client = OpenAI(api_key=api_key)

            sys = (
                "Ești un asistent care completează formulare agricole din R. Moldova. "
                "Întoarce STRICT un JSON object cu chei EXACT câmpurile primite."
            )
            user_msg = {
                "language": language,
                "instructions": instructions or "Completează realist, concis și formal.",
                "profile": profile or {},
                "fields": [{"id": f["id"], "label": f["label"], "context": f["context"]} for f in fields],
            }
            resp = client.chat.completions.create(
                model="GPT_MODEL",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": sys},
                    {"role": "user", "content": json.dumps(user_msg, ensure_ascii=False)},
                    {"role": "user", "content": "Returnează DOAR JSON-ul cu {id: valoare}."},
                ],
                temperature=0.2,
            )
            raw = resp.choices[0].message.content
            data = json.loads(raw)
            out = {}
            want = {f["id"] for f in fields}
            for k, v in (data or {}).items():
                if k in want:
                    out[k] = str(v)
            for f in fields:
                out.setdefault(f["id"], f"[completați: {f['label']}]")
            return out
    except Exception:
        pass

    out = {}
    profile = profile or {}
    for f in fields:
        label = f["label"]
        key_guess = _slug(label)
        v = profile.get(key_guess) or profile.get(label) or f"[exemplu automat pentru: {label}]"
        out[f["id"]] = str(v)
    return out

def _download_bytes(url: str) -> bytes:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    return r.content

def _generate_pdf_for_profile(profile: dict) -> io.BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Farm Profile", ln=True)
    
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Farmer Info", ln=True)
    pdf.set_font("Arial", size=12)
    user = profile.get("user", {})
    if user:
        for k, v in user.items():
            pdf.multi_cell(0, 6, f"{k}: {v}")
    else:
        pdf.cell(0, 6, "No user info available", ln=True)
    
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Animals", ln=True)
    pdf.set_font("Arial", size=12)
    animals = profile.get("animals", [])
    if animals:
        for a in animals:
            pdf.multi_cell(0, 6, ", ".join(f"{k}: {v}" for k, v in a.items()))
    else:
        pdf.cell(0, 6, "No animals available", ln=True)
    
    # Similarly for fields and vehicles
    for section_name, items in [("Fields", profile.get("fields", [])),
                                ("Vehicles", profile.get("vehicles", []))]:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, section_name, ln=True)
        pdf.set_font("Arial", size=12)
        if items:
            for item in items:
                pdf.multi_cell(0, 6, ", ".join(f"{k}: {v}" for k, v in item.items()))
        else:
            pdf.cell(0, 6, f"No {section_name.lower()} available", ln=True)
    
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    bio = io.BytesIO(pdf_bytes)
    bio.seek(0)
    return bio
    
# mongo db profile loader

def _load_profile_by_owner_id(owner_id: str) -> dict:
    """
    Load farm data first, then user.
    """
    try:
        owner_id = owner_id.strip()

        farm_db = client["FarmXpertDB"]

        animals = list(farm_db.animals.find({"OwnerId": owner_id}))
        fields = list(farm_db.fields.find({"OwnerId": owner_id}))
        vehicles = list(farm_db.vehicles.find({"OwnerId": owner_id}))

        print("Owner ID:", owner_id)
        print("Animals:", animals)
        print("Fields:", fields)
        print("Vehicles:", vehicles)

        # Now attempt to fetch user
        try:
            user_doc = client["default"].users.find_one({"_id": ObjectId(owner_id)})
        except Exception:
            user_doc = {}

        return {
            "user": user_doc or {},
            "animals": animals,
            "fields": fields,
            "vehicles": vehicles,
        }

    except Exception as e:
        print("Error loading profile:", e)
        return {}
# api endpoints

@scraper_bp.post("/complete-docx")
def complete_docx():
    """
    Autocomplete DOCX fields using user profile from MongoDB.
    ---
    tags:
      - Document Processing
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          properties:
            url:
              type: string
              description: Direct URL to DOCX file
            ownerId:
              type: string
              description: MongoDB user ID
            instructions:
              type: string
              description: Optional AI instructions for field completion
            language:
              type: string
              default: "ro"
              description: Language code (default Romanian)
            filename:
              type: string
              default: "document_autocomplete.docx"
              description: Output filename
    responses:
      200:
        description: Completed DOCX file
        schema:
          type: file
      400:
        description: Missing or invalid parameters
      500:
        description: Processing error
      502:
        description: Failed to download source document
    """
    data = request.get_json(silent=True) or {}
    
    url = (data.get("url") or "").strip()
    owner_id = (data.get("ownerId") or "").strip()
    filename = (data.get("filename") or "document_autocomplete.docx").strip()
    language = (data.get("language") or "ro").strip().lower()
    instructions = data.get("instructions")

    # Validate inputs
    if not url:
        return jsonify(error="missing_url", details="Provide 'url' pointing to a .docx file."), 400
    
    if not url.lower().endswith(".docx"):
        return jsonify(error="unsupported   _format", details="Only .docx files are supported."), 400

    # Load profile from MongoDB if ownerId provided
    profile = {}
    if owner_id:
        try:
            profile = _load_profile_by_owner_id(owner_id)
        except Exception as e:
            return jsonify(error="profile_load_failed", details=str(e)), 500

    try:
        # 1) Download source DOCX
        content = _download_bytes(url)
        doc = Document(io.BytesIO(content))

        # 2) Extract blank fields
        fields = _extract_doc_fields(doc)

        # 3) Get AI suggestions
        ai_suggestions = _ai_fill_fields(fields, profile=profile, instructions=instructions, language=language)

        # 4) Flatten profile for label matching
        profile_flat = {k.lower(): v for k, v in _flatten_profile(profile).items()}
        
        # Merge AI suggestions with profile data
        suggestions = {}
        suggestions.update(profile_flat)
        suggestions.update({k.lower().replace("_", " "): v for k, v in ai_suggestions.items()})

        # 5) Apply suggestions to document
        _apply_suggestions_inline(doc, suggestions)
        _apply_label_and_cell_fill(doc, suggestions)

        # 6) Add header explaining markers
        hdr_text = "«Aceste marcaje indică sugestiile auto-completate. Editați la nevoie înainte de depunere.»"
        try:
            if doc.paragraphs:
                doc.paragraphs[0].insert_paragraph_before(hdr_text)
            else:
                doc.add_paragraph(hdr_text)
        except Exception:
            try:
                doc.add_paragraph(hdr_text)
            except Exception:
                pass

        # 7) Return completed document
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return send_file(
            bio,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )

    except requests.HTTPError as e:
        return jsonify(error="download_failed", details=str(e)), 502
    except Exception as e:
        return jsonify(error="autocomplete_failed", details=str(e)), 500



@scraper_bp.post("/extract-profile")
def extract_profile():
    """
    Extract user profile and farm data from MongoDB.
    ---
    tags:
      - Profile Management
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          properties:
            ownerId:
              type: string
              description: MongoDB user ID (required)
            callbackAPI:
              type: string
              description: Optional URL to POST results to
    responses:
      200:
        description: User profile with farm data
        schema:
          type: object
          properties:
            OwnerId:
              type: string
            user:
              type: object
            animals:
              type: array
              items:
                type: object
            fields:
              type: array
              items:
                type: object
            vehicles:
              type: array
              items:
                type: object
            count:
              type: object
              properties:
                animals:
                  type: integer
                fields:
                  type: integer
                vehicles:
                  type: integer
      400:
        description: Missing ownerId
      500:
        description: Extraction failed
    """
    data = request.get_json(silent=True) or {}
    
    owner_id = (data.get("OwnerId") or "").strip()
    callback_api = (data.get("callbackAPI") or "").strip()

    if not owner_id:
        return jsonify(error="missing_owner_id", details="Provide 'ownerId' field."), 400

    try:
        # Load all data for this owner
        profile = _load_profile_by_owner_id(owner_id)
        
        result = {
            "OwnerId": owner_id,
            "user": profile.get("user", {}),
            "animals": profile.get("animals", []),
            "fields": profile.get("fields", []),
            "vehicles": profile.get("vehicles", []),
            "count": {
                "animals": len(profile.get("animals", [])),
                "fields": len(profile.get("fields", [])),
                "vehicles": len(profile.get("vehicles", []))
            }
        }

        # If callback API provided, forward the data
        if callback_api:
            try:
                resp = requests.post(callback_api, json=result, timeout=10)
                return jsonify({
                    "message": "Data sent to callback API",
                    "callback_status": resp.status_code,
                    "data": result
                })
            except Exception as e:
                return jsonify({
                    "error": "callback_failed",
                    "details": str(e),
                    "data": result
                }), 500

        return jsonify(result)

    except Exception as e:
        return jsonify(error="extraction_failed", details=str(e)), 500
    
@scraper_bp.post("/doc-gen")
def doc_gen():
    """
    Generate a PDF document for a farmer profile.
    ---
    tags:
      - Document Generation
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          properties:
            ownerId:
              type: string
              description: MongoDB user ID (required)
            callbackAPI:
              type: string
              description: Optional URL to POST the generated PDF to
    responses:
      200:
        description: Generated PDF file or confirmation of callback delivery
        schema:
          type: object
          properties:
            message:
              type: string
            callback_status:
              type: integer
      400:
        description: Missing ownerId
      500:
        description: PDF generation failed
    """
    data = request.get_json(silent=True) or {}
    owner_id = (data.get("OwnerId") or "").strip()
    callback_api = (data.get("callbackAPI") or "").strip()

    if not owner_id:
        return jsonify(error="missing_owner_id", details="Provide 'ownerId'."), 400

    try:
        # Load farmer profile from MongoDB
        profile = _load_profile_by_owner_id(owner_id)

        # Generate PDF
        doc_bytes = _generate_pdf_for_profile(profile)
        filename = f"{owner_id}_farm_profile.pdf"
        mimetype = "application/pdf"

        # If callbackAPI exists, POST it there
        if callback_api:
            files = {"processedData": (filename, doc_bytes.getvalue())}
            resp = requests.post(callback_api, data={"ownerId": owner_id}, files=files, timeout=20)
            return jsonify({
                "message": "PDF sent to callback API",
                "callback_status": resp.status_code
            })

        # Otherwise return PDF directly
        return send_file(
            doc_bytes,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify(error="doc_generation_failed", details=str(e)), 500


# Register blueprints
bp.register_blueprint(scraper_bp)